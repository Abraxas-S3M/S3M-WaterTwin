"""Tests for the customer-document text-extraction parser (read-only).

Covers, per the phase requirements:

* each supported format (.pdf / .docx / .md / .txt) extracts with correct,
  resolvable chunk character offsets;
* a scanned / image-only PDF (no text layer) reports a clear ``unparsed``
  reason and never crashes;
* a macro-enabled ``.docm`` is rejected;
* a malformed PDF reports ``parse_failed`` (no crash, no hang);
* chunk boundaries are deterministic (same bytes -> same chunks); and
* the parser never executes active content (it only reads a text layer).
"""

from __future__ import annotations

from io import BytesIO

from app.parsers.document import (
    ParseStatus,
    chunk_text,
    parse_document,
)


# --------------------------------------------------------------------------- #
# Minimal in-memory document builders (synthetic; no fixtures on disk)
# --------------------------------------------------------------------------- #


def _obj(num: int, body: bytes) -> bytes:
    return f"{num} 0 obj\n".encode() + body + b"\nendobj\n"


def _assemble_pdf(objects: dict[int, bytes]) -> bytes:
    out = BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for num in sorted(objects):
        offsets[num] = out.tell()
        out.write(_obj(num, objects[num]))
    xref_pos = out.tell()
    size = max(objects) + 1
    out.write(f"xref\n0 {size}\n".encode())
    out.write(b"0000000000 65535 f \n")
    for num in range(1, size):
        if num in offsets:
            out.write(f"{offsets[num]:010d} 00000 n \n".encode())
        else:
            out.write(b"0000000000 65535 f \n")
    out.write(
        f"trailer\n<< /Size {size} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF".encode()
    )
    return out.getvalue()


def make_text_pdf(pages: list[list[str]]) -> bytes:
    """Build a PDF with a real text layer (one Tj per line)."""
    objects: dict[int, bytes] = {}
    page_nums = [4 + 2 * i for i in range(len(pages))]
    content_nums = [5 + 2 * i for i in range(len(pages))]
    kids = " ".join(f"{p} 0 R" for p in page_nums)
    objects[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    objects[2] = f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode()
    objects[3] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
    for i, lines in enumerate(pages):
        objects[page_nums[i]] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {content_nums[i]} 0 R >>"
        ).encode()
        stream = "BT\n/F1 12 Tf\n72 720 Td\n"
        for j, line in enumerate(lines):
            if j > 0:
                stream += "0 -14 Td\n"
            esc = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            stream += f"({esc}) Tj\n"
        stream += "ET"
        sb = stream.encode()
        objects[content_nums[i]] = (
            f"<< /Length {len(sb)} >>\nstream\n".encode() + sb + b"\nendstream"
        )
    return _assemble_pdf(objects)


def make_image_only_pdf() -> bytes:
    """Build a single-page PDF whose content draws a rectangle -- no text layer."""
    objects: dict[int, bytes] = {}
    objects[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    objects[2] = b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>"
    objects[3] = (
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>"
    )
    stream = b"100 100 200 200 re f"
    objects[4] = b"<< /Length %d >>\nstream\n" % len(stream) + stream + b"\nendstream"
    return _assemble_pdf(objects)


def make_docx(paragraphs: list[tuple[str, str]]) -> bytes:
    """Build a .docx from ``(style, text)`` paragraphs (style '' = Normal)."""
    from docx import Document

    document = Document()
    for style, text in paragraphs:
        if style:
            document.add_paragraph(text, style=style)
        else:
            document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# Offset invariant helper
# --------------------------------------------------------------------------- #


def _assert_offsets_resolve(parsed) -> None:
    """Every chunk's offsets must resolve back to its text in the source."""
    assert parsed.chunks
    for chunk in parsed.chunks:
        assert parsed.text[chunk.char_start : chunk.char_end] == chunk.text
        assert chunk.char_start < chunk.char_end
        assert chunk.source_document_id == parsed.source_document_id


# --------------------------------------------------------------------------- #
# Markdown / text
# --------------------------------------------------------------------------- #


def test_markdown_extracts_with_sections_and_resolvable_offsets():
    body = (
        "# Cleaning Procedure\n\n"
        "Isolate the membrane train before starting.\n\n"
        "## Chemical Dose\n\n"
        "Dose citric acid to pH 4 and recirculate for thirty minutes.\n"
    )
    parsed = parse_document("sop.md", body.encode(), source_document_id="doc-md")
    assert parsed.status is ParseStatus.parsed
    assert parsed.media_type == "text/markdown"
    _assert_offsets_resolve(parsed)
    # Sections are attributed from the nearest preceding heading.
    sections = {c.section for c in parsed.chunks}
    assert "Cleaning Procedure" in sections
    assert "Chemical Dose" in sections
    dose_chunk = next(c for c in parsed.chunks if "citric acid" in c.text)
    assert dose_chunk.section == "Chemical Dose"


def test_plaintext_extracts_and_offsets_resolve():
    body = "Line one of the log.\n\nLine two of the log.\n"
    parsed = parse_document("notes.txt", body.encode(), source_document_id="doc-txt")
    assert parsed.status is ParseStatus.parsed
    assert parsed.media_type == "text/plain"
    _assert_offsets_resolve(parsed)


def test_empty_text_document_is_unparsed():
    parsed = parse_document("blank.txt", b"   \n\n  \n", source_document_id="doc-empty")
    assert parsed.status is ParseStatus.unparsed
    assert parsed.reason and "no extractable text" in parsed.reason
    assert parsed.chunks == []


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #


def test_pdf_text_layer_extracts_with_pages_and_offsets():
    data = make_text_pdf(
        [
            ["High pressure pump isolation", "Lock out the breaker first"],
            ["Membrane cleaning schedule and dosing"],
        ]
    )
    parsed = parse_document("manual.pdf", data, source_document_id="doc-pdf")
    assert parsed.status is ParseStatus.parsed
    assert parsed.page_count == 2
    assert parsed.media_type == "application/pdf"
    _assert_offsets_resolve(parsed)
    # Chunks carry the page they came from and never cross a page boundary.
    pages = {c.page for c in parsed.chunks}
    assert pages == {1, 2}
    p1 = next(c for c in parsed.chunks if c.page == 1)
    assert "isolation" in p1.text.lower()


def test_image_only_pdf_is_reported_unparsed_without_crash():
    data = make_image_only_pdf()
    parsed = parse_document("scan.pdf", data, source_document_id="doc-scan")
    assert parsed.status is ParseStatus.unparsed
    assert parsed.reason and "image-only" in parsed.reason
    assert "OCR is not performed" in parsed.reason
    assert parsed.chunks == []


def test_malformed_pdf_reports_parse_failed_without_crash():
    parsed = parse_document(
        "broken.pdf", b"%PDF-1.4 this is not a real pdf body", source_document_id="doc-bad"
    )
    assert parsed.status is ParseStatus.parse_failed
    assert parsed.reason and "could not parse" in parsed.reason
    assert parsed.chunks == []


def test_truncated_pdf_reports_parse_failed_without_crash():
    good = make_text_pdf([["some text here"]])
    parsed = parse_document("trunc.pdf", good[: len(good) // 2], source_document_id="doc-tr")
    assert parsed.status in {ParseStatus.parse_failed, ParseStatus.unparsed}
    # Either outcome is acceptable; the invariant is that it did not raise.


# --------------------------------------------------------------------------- #
# DOCX / DOCM
# --------------------------------------------------------------------------- #


def test_docx_extracts_paragraphs_and_headings():
    data = make_docx(
        [
            ("Heading 1", "Startup Procedure"),
            ("", "Confirm the feed valve is open before energising the pump."),
            ("Heading 2", "Shutdown"),
            ("", "De-energise and lock out the high pressure pump."),
        ]
    )
    parsed = parse_document("sop.docx", data, source_document_id="doc-docx")
    assert parsed.status is ParseStatus.parsed
    assert "wordprocessingml" in parsed.media_type
    _assert_offsets_resolve(parsed)
    sections = {c.section for c in parsed.chunks}
    assert "Startup Procedure" in sections
    assert "Shutdown" in sections


def test_docm_is_rejected():
    # Content is irrelevant: a .docm is refused on its extension (macro risk).
    parsed = parse_document("with_macros.docm", b"PK\x03\x04 anything", source_document_id="d")
    assert parsed.status is ParseStatus.rejected
    assert parsed.reason and ".docm" in parsed.reason
    assert parsed.chunks == []


def test_other_macro_formats_are_rejected():
    for name in ("book.xlsm", "deck.pptm", "template.dotm"):
        parsed = parse_document(name, b"PK\x03\x04", source_document_id="d")
        assert parsed.status is ParseStatus.rejected


def test_unsupported_extension_is_parse_failed_not_a_crash():
    parsed = parse_document("drawing.dwg", b"\x00\x01\x02", source_document_id="d")
    assert parsed.status is ParseStatus.parse_failed
    assert parsed.reason and "unsupported file type" in parsed.reason


# --------------------------------------------------------------------------- #
# Determinism + chunking behaviour
# --------------------------------------------------------------------------- #


def test_chunking_is_deterministic():
    body = "\n\n".join(f"Paragraph {i} with several words in it." for i in range(20))
    first = parse_document("a.md", body.encode(), source_document_id="doc-det")
    second = parse_document("a.md", body.encode(), source_document_id="doc-det")
    assert [c.as_dict() for c in first.chunks] == [c.as_dict() for c in second.chunks]
    assert first.sha256 == second.sha256


def test_long_paragraph_is_split_into_bounded_windows():
    long_para = "word " * 1000  # ~5000 chars, single paragraph
    chunks = chunk_text(long_para, source_document_id="doc-long", max_chars=500)
    assert len(chunks) > 1
    for c in chunks:
        assert c.char_end - c.char_start <= 500
    # Windows tile the paragraph contiguously.
    assert chunks[0].char_start == 0
    for prev, nxt in zip(chunks, chunks[1:]):
        assert nxt.char_start == prev.char_end


def test_sha256_is_content_addressed():
    a = parse_document("x.txt", b"hello", source_document_id="a")
    b = parse_document("y.txt", b"hello", source_document_id="b")
    assert a.sha256 == b.sha256  # same bytes -> same hash regardless of name
