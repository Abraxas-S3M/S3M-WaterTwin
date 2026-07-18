"""Customer-document text extraction + deterministic chunking (read-only).

This module extracts the **text layer** of a customer-supplied document and
splits it into stable, deterministic chunks. It is intentionally defensive: it
never executes active content and never reaches out over the network.

Supported inputs
----------------
* ``.pdf``  -- text layer only, via :mod:`pypdf`. No JavaScript is executed, no
  embedded files are extracted, and no form-field values are evaluated. A PDF
  with no recoverable text layer (a scanned / image-only document) is reported
  as :data:`ParseStatus.unparsed` with a clear reason -- **no OCR is performed**.
* ``.docx`` -- via :mod:`python-docx`. Macros are never executed.
* ``.md`` / ``.txt`` -- read directly as UTF-8 text.

Macro-enabled Office formats (``.docm`` and friends) are **rejected** outright.
Malformed input is reported as :data:`ParseStatus.parse_failed`; the parser
never raises to its caller and never hangs.

Every :class:`Chunk` records its source document id, its page (PDF) or section
(Markdown / Word heading), and the character offsets ``[char_start, char_end)``
into the extracted text -- so a downstream citation can resolve to a real
location in a real file. The chunk boundaries are deterministic: the same bytes
always produce the same chunks.
"""

from __future__ import annotations

import hashlib
import os
import re
from dataclasses import dataclass, field
from enum import Enum
from io import BytesIO
from typing import Optional

#: File extensions this parser can extract text from.
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".docx", ".md", ".txt"})

#: Macro-enabled / active-content Office formats that are rejected on sight. The
#: task requires rejecting ``.docm``; the rest are refused for the same reason
#: (they can carry executable macros) so no active content is ever opened.
_MACRO_EXTENSIONS: frozenset[str] = frozenset({".docm", ".dotm", ".xlsm", ".pptm", ".xlsb"})

#: Default maximum characters per chunk. Chunk boundaries prefer paragraph
#: breaks; an over-long paragraph is split into fixed windows of this size.
DEFAULT_MAX_CHARS = 1200

#: Markdown/setext-style heading line (``# Title`` .. ``###### Title``).
_HEADING_RE = re.compile(r"^[ \t]*(#{1,6})[ \t]+(.+?)[ \t]*#*[ \t]*$", re.MULTILINE)

#: A block of one or more consecutive non-blank lines (a "paragraph").
_PARAGRAPH_RE = re.compile(r"(?:^[ \t]*\S.*(?:\n|$))+", re.MULTILINE)


class ParseStatus(str, Enum):
    """Outcome of a parse attempt."""

    #: Text was extracted and chunked successfully.
    parsed = "parsed"
    #: The format was recognised but no usable text could be recovered (e.g. a
    #: scanned / image-only PDF, or an empty document). Never OCR'd here.
    unparsed = "unparsed"
    #: The input was malformed / could not be decoded. No crash, no hang.
    parse_failed = "parse_failed"
    #: The input was refused for policy reasons (macro-enabled Office document).
    rejected = "rejected"


@dataclass(frozen=True)
class Chunk:
    """A deterministic text chunk with a resolvable source location.

    ``char_start`` / ``char_end`` are offsets into :attr:`ParsedDocument.text`;
    ``text == parsed.text[char_start:char_end]`` always holds. ``page`` is set
    for PDF sources (1-based); ``section`` is set when a heading precedes the
    chunk (Markdown headings or Word heading styles).
    """

    chunk_id: str
    source_document_id: str
    text: str
    char_start: int
    char_end: int
    page: Optional[int] = None
    section: Optional[str] = None

    def as_dict(self) -> dict:
        return {
            "chunk_id": self.chunk_id,
            "source_document_id": self.source_document_id,
            "text": self.text,
            "char_start": self.char_start,
            "char_end": self.char_end,
            "page": self.page,
            "section": self.section,
        }


@dataclass
class ParsedDocument:
    """The result of parsing a single customer document."""

    source_document_id: str
    filename: str
    media_type: str
    sha256: str
    status: ParseStatus
    text: str = ""
    reason: Optional[str] = None
    page_count: Optional[int] = None
    chunks: list[Chunk] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        """True when text was extracted and at least one chunk was produced."""
        return self.status is ParseStatus.parsed and bool(self.chunks)

    def as_store_chunks(self) -> list[dict]:
        """Chunks as plain dicts for the DocumentStore ingest contract."""
        return [c.as_dict() for c in self.chunks]


def sha256_hex(data: bytes) -> str:
    """Return the hex SHA-256 of ``data`` (content integrity / dedup key)."""
    return hashlib.sha256(data).hexdigest()


def _extension(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def _media_type(ext: str) -> str:
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".md": "text/markdown",
        ".txt": "text/plain",
    }.get(ext, "application/octet-stream")


def _markdown_sections(text: str) -> list[tuple[int, str]]:
    """Return ``(offset, title)`` for each Markdown heading, in document order."""
    return [(m.start(), m.group(2).strip()) for m in _HEADING_RE.finditer(text)]


def _section_for(offset: int, sections: list[tuple[int, str]]) -> Optional[str]:
    """The title of the nearest heading at or before ``offset`` (or ``None``)."""
    current: Optional[str] = None
    for start, title in sections:
        if start <= offset:
            current = title
        else:
            break
    return current


def _section_regions(
    text: str, sections: list[tuple[int, str]]
) -> list[tuple[Optional[int], int, int]]:
    """Split ``text`` into ``(page=None, start, end)`` regions at heading offsets.

    Each heading begins a new region so a chunk never spans two sections; text
    before the first heading is its own region.
    """
    boundaries = sorted({0, len(text)} | {offset for offset, _ in sections})
    regions: list[tuple[Optional[int], int, int]] = []
    for start, end in zip(boundaries, boundaries[1:]):
        if end > start:
            regions.append((None, start, end))
    return regions or [(None, 0, len(text))]


def _paragraph_spans(text: str, start: int, end: int) -> list[tuple[int, int]]:
    """Absolute ``(start, end)`` spans of paragraphs within ``text[start:end]``."""
    region = text[start:end]
    spans: list[tuple[int, int]] = []
    for match in _PARAGRAPH_RE.finditer(region):
        block_start = start + match.start()
        block_end = start + match.end()
        # Trim trailing whitespace/newline so text[cs:ce] is the clean block.
        while block_end > block_start and text[block_end - 1].isspace():
            block_end -= 1
        if block_end > block_start:
            spans.append((block_start, block_end))
    return spans


def _emit_chunks(
    *,
    text: str,
    region_start: int,
    region_end: int,
    page: Optional[int],
    sections: list[tuple[int, str]],
    source_document_id: str,
    start_index: int,
    max_chars: int,
) -> list[Chunk]:
    """Chunk ``text[region_start:region_end]`` with paragraph-preferred, stable
    boundaries. Offsets are absolute into ``text``."""
    chunks: list[Chunk] = []
    index = start_index
    cur_start: Optional[int] = None
    cur_end: Optional[int] = None

    def flush() -> None:
        nonlocal cur_start, cur_end, index
        if cur_start is None or cur_end is None:
            return
        chunks.append(
            Chunk(
                chunk_id=f"{source_document_id}:c{index}",
                source_document_id=source_document_id,
                text=text[cur_start:cur_end],
                char_start=cur_start,
                char_end=cur_end,
                page=page,
                section=_section_for(cur_start, sections),
            )
        )
        index += 1
        cur_start = None
        cur_end = None

    for p_start, p_end in _paragraph_spans(text, region_start, region_end):
        if p_end - p_start > max_chars:
            # Over-long paragraph: flush any pending chunk, then hard-split it
            # into deterministic fixed windows.
            flush()
            window = p_start
            while window < p_end:
                w_end = min(window + max_chars, p_end)
                chunks.append(
                    Chunk(
                        chunk_id=f"{source_document_id}:c{index}",
                        source_document_id=source_document_id,
                        text=text[window:w_end],
                        char_start=window,
                        char_end=w_end,
                        page=page,
                        section=_section_for(window, sections),
                    )
                )
                index += 1
                window = w_end
            continue
        if cur_start is None:
            cur_start, cur_end = p_start, p_end
        elif p_end - cur_start <= max_chars:
            cur_end = p_end
        else:
            flush()
            cur_start, cur_end = p_start, p_end
    flush()
    return chunks


def chunk_text(
    text: str,
    *,
    source_document_id: str,
    pages: Optional[list[tuple[int, int, int]]] = None,
    sections: Optional[list[tuple[int, str]]] = None,
    max_chars: int = DEFAULT_MAX_CHARS,
) -> list[Chunk]:
    """Split ``text`` into deterministic chunks.

    ``pages`` (PDF) is a list of ``(page_number, char_start, char_end)`` regions;
    when supplied, chunks never cross a page boundary and each carries its page.
    ``sections`` is a list of ``(offset, title)`` used to tag each chunk with the
    heading it falls under and to force a chunk boundary at each heading. When
    ``pages`` is omitted the text is split at section boundaries (or treated as a
    single region when it has no headings).
    """
    if sections is None:
        sections = _markdown_sections(text)
    if pages:
        regions: list[tuple[Optional[int], int, int]] = pages
    else:
        regions = _section_regions(text, sections)
    chunks: list[Chunk] = []
    for page, region_start, region_end in regions:
        chunks.extend(
            _emit_chunks(
                text=text,
                region_start=region_start,
                region_end=region_end,
                page=page,
                sections=sections,
                source_document_id=source_document_id,
                start_index=len(chunks),
                max_chars=max_chars,
            )
        )
    return chunks


# --------------------------------------------------------------------------- #
# Format-specific text extraction
# --------------------------------------------------------------------------- #


def _extract_pdf(data: bytes) -> tuple[str, list[tuple[int, int, int]], int]:
    """Extract the PDF *text layer* only. Returns ``(text, pages, page_count)``.

    Text-layer extraction only: no JavaScript is executed, no embedded files are
    read, and no form fields are evaluated (pypdf's ``extract_text`` reads page
    content streams exclusively).
    """
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    pages: list[tuple[int, int, int]] = []
    cursor = 0
    for number, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if number > 1:
            # Deterministic page separator so offsets stay stable + resolvable.
            parts.append("\n\n")
            cursor += 2
        start = cursor
        parts.append(page_text)
        cursor += len(page_text)
        pages.append((number, start, cursor))
    return "".join(parts), pages, len(pages)


def _extract_docx(data: bytes) -> tuple[str, list[tuple[int, str]]]:
    """Extract Word text + heading sections. Macros are never executed."""
    from docx import Document

    document = Document(BytesIO(data))
    parts: list[str] = []
    sections: list[tuple[int, str]] = []
    cursor = 0
    for para in document.paragraphs:
        line = para.text
        style = getattr(getattr(para, "style", None), "name", "") or ""
        if line.strip() and (style.startswith("Heading") or style == "Title"):
            sections.append((cursor, line.strip()))
        if parts:
            parts.append("\n")
            cursor += 1
        parts.append(line)
        cursor += len(line)
    return "".join(parts), sections


def _decode_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


# --------------------------------------------------------------------------- #
# Public entry point
# --------------------------------------------------------------------------- #


def parse_document(
    filename: str,
    data: bytes,
    *,
    source_document_id: Optional[str] = None,
    max_chars: int = DEFAULT_MAX_CHARS,
) -> ParsedDocument:
    """Parse ``data`` (the bytes of ``filename``) into a :class:`ParsedDocument`.

    Never raises: any extraction failure is reported as
    :data:`ParseStatus.parse_failed`; macro-enabled Office documents are
    :data:`ParseStatus.rejected`; a text-less PDF is :data:`ParseStatus.unparsed`.
    """
    ext = _extension(filename)
    sha = sha256_hex(data)
    doc_id = source_document_id or f"ingest-{sha[:16]}"
    media_type = _media_type(ext)

    def result(status: ParseStatus, **kw) -> ParsedDocument:
        return ParsedDocument(
            source_document_id=doc_id,
            filename=filename,
            media_type=media_type,
            sha256=sha,
            status=status,
            **kw,
        )

    if ext in _MACRO_EXTENSIONS:
        return result(
            ParseStatus.rejected,
            reason=(
                f"macro-enabled Office documents ({ext}) are not accepted; "
                "convert to .docx or .pdf and re-upload"
            ),
        )
    if ext not in SUPPORTED_EXTENSIONS:
        return result(
            ParseStatus.parse_failed,
            reason=f"unsupported file type '{ext or '(none)'}'",
        )

    try:
        if ext == ".pdf":
            text, pages, page_count = _extract_pdf(data)
            if not text.strip():
                return result(
                    ParseStatus.unparsed,
                    page_count=page_count,
                    reason=(
                        "no extractable text layer (likely a scanned / image-only "
                        "PDF); OCR is not performed"
                    ),
                )
            chunks = chunk_text(
                text, source_document_id=doc_id, pages=pages, max_chars=max_chars
            )
            return result(
                ParseStatus.parsed, text=text, page_count=page_count, chunks=chunks
            )

        if ext == ".docx":
            text, sections = _extract_docx(data)
            if not text.strip():
                return result(
                    ParseStatus.unparsed,
                    reason="document contains no extractable text",
                )
            chunks = chunk_text(
                text, source_document_id=doc_id, sections=sections, max_chars=max_chars
            )
            return result(ParseStatus.parsed, text=text, chunks=chunks)

        # .md / .txt
        text = _decode_text(data)
        if not text.strip():
            return result(
                ParseStatus.unparsed, reason="document contains no extractable text"
            )
        chunks = chunk_text(text, source_document_id=doc_id, max_chars=max_chars)
        return result(ParseStatus.parsed, text=text, chunks=chunks)
    except Exception as exc:  # malformed input -> reported, never raised
        return result(
            ParseStatus.parse_failed,
            reason=f"could not parse {ext} document: {type(exc).__name__}",
        )
